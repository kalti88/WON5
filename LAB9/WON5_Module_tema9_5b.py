import statistics
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def import_data_create_catalog(filename: str):
    """Importa datele pentru crearea catalogului cu medii

    Parameters
    ----------
    filename: str
                Fisierul ce contine numele si notele elevilor.

    Returns
    -------
    dict:
            Returneaza catalogul elevilor cu mediile calculate.
    """
    catalog = {}
    with open(filename, 'r') as file:
        for line in file:
            ln = line.strip().split(';')
            note = ln.copy()
            note.remove(note[0])
            if note[0]:
                note = [int(i) for i in note]
                catalog[ln[0]] = round(statistics.mean(note), 2)
    return catalog


def premianti_3(c: dict):
    """Creeaza lista cu cei 3 premianti.

    Parameters
    ----------
    c : dict
        Dictionar care contine catalogul clasei

    Returns
    -------
    list
        "Lista cu cei 3 premianti.
    """
    lista_p = []
    for e in c:
        lista_p.append((c[e], e))
    lista_p.sort(reverse=True)
    return lista_p[:3]


def create_diploma(filename: str, premiantii, counter):
    """Importa datele pentru crearea diplomelor pentru premianti

    Parameters
    ----------
    filename: str
                Fisierul ce contine numele si notele elevilor.
    premiantii: list
                Lista cu premiantii
    counter: int
                Numar intreg pentru ordinea diplomelor

    Returns
    -------
    None
    """
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(15)

    document.add_heading('Diploma', 0)
    document.add_paragraph()

    with open(filename, 'r') as file:
        i = 1
        for line in file:
            if i == 1:
                p = document.add_paragraph((line.replace('{}', ' ').strip(), ' '))
                p.add_run(str(counter)).bold = True
            elif i == 2:
                document.add_paragraph(((line.replace('{}', ':')).strip(), ' '))
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(str(premiantii[counter-1][1])).bold = True
            else:
                p = document.add_paragraph((line.replace('{}', ' ').strip(), ' '))
                p.add_run(str(premiantii[counter-1][0])).bold = True
            document.save(f'diploma-{counter}.docx')
            i += 1
